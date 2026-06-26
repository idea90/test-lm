from fastapi import FastAPI, Request, status, File, UploadFile, Form
from fastapi.responses import JSONResponse, FileResponse, StreamingResponse
from starlette.middleware.sessions import SessionMiddleware
import os
import pypdf
import io
from dotenv import load_dotenv
from werkzeug.security import generate_password_hash, check_password_hash
import pytesseract
from pdf2image import convert_from_bytes
import database
import llm_helper
import gemini_helper

load_dotenv()

app = FastAPI()

app.add_middleware(
    SessionMiddleware, 
    secret_key=os.getenv("SECRET_KEY", "test-lm-default-secret-key-123456")
)

# Initialize SQLite database
database.init_db()

POPPLER_PATH = os.getenv("POPPLER_PATH") or r"C:\Users\idea\AppData\Local\Microsoft\WinGet\Packages\oschwartz10612.Poppler_Microsoft.Winget.Source_8wekyb3d8bbwe\poppler-25.07.0\Library\bin"
pytesseract.pytesseract.tesseract_cmd = os.getenv("TESSERACT_CMD") or r"C:\Program Files\Tesseract-OCR\tesseract.exe"

UPLOAD_AVATAR_FOLDER = os.path.join(os.path.dirname(__file__), 'uploads', 'avatars')
os.makedirs(UPLOAD_AVATAR_FOLDER, exist_ok=True)

# Helper to check authentication
def get_current_user(request: Request):
    user_id = request.session.get('user_id')
    if not user_id:
        return None
    return user_id

def extract_text_from_pdf(file_stream, page_start=1, page_end=None, force_ocr=False, exclude_pages=None):
    try:
        from concurrent.futures import ThreadPoolExecutor
        
        if exclude_pages is None:
            exclude_pages = set()
            
        reader = pypdf.PdfReader(file_stream)
        total_pages = len(reader.pages)
        
        start_idx = max(0, (page_start or 1) - 1)
        end_idx = min(total_pages, page_end if page_end is not None else total_pages)
        
        file_stream.seek(0)
        pdf_bytes = file_stream.read()
        
        results = [None] * (end_idx - start_idx)
        pages_to_ocr = []
        
        for idx, i in enumerate(range(start_idx, end_idx)):
            page_num = i + 1
            if page_num in exclude_pages:
                results[idx] = ""
                continue
                
            page_text = ""
            if not force_ocr:
                page = reader.pages[i]
                page_text = page.extract_text() or ""
            
            if not force_ocr and page_text and len(page_text.strip()) >= 15:
                results[idx] = page_text
            else:
                pages_to_ocr.append((idx, page_num))
                
        if pages_to_ocr:
            print(f"Running parallel Tesseract OCR on pages: {[p[1] for p in pages_to_ocr]}...")
            
            def ocr_page(page_num):
                try:
                    images = convert_from_bytes(pdf_bytes, first_page=page_num, last_page=page_num, poppler_path=POPPLER_PATH)
                    if images:
                        ocr_text = pytesseract.image_to_string(images[0], lang='lao+eng')
                        if ocr_text and len(ocr_text.strip()) > 5:
                            return page_num, ocr_text.strip()
                except Exception as ocr_err:
                    print(f"OCR failed for page {page_num}: {ocr_err}")
                return page_num, ""
            
            max_workers = min(len(pages_to_ocr), 4, os.cpu_count() or 4)
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                futures = {
                    executor.submit(ocr_page, page_num): idx
                    for idx, page_num in pages_to_ocr
                }
                for future in futures:
                    idx = futures[future]
                    page_num, ocr_result = future.result()
                    if ocr_result:
                        results[idx] = ocr_result
                        print(f"Successfully OCR'd page {page_num} using parallel Tesseract.")
                    else:
                        results[idx] = ""
                        
        text = "\n".join([r for r in results if r])
        return text.strip()
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return None

def extract_text_from_docx(file_stream):
    try:
        from docx import Document
        doc = Document(file_stream)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text.strip()
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return None

def extract_text_from_image(file_stream):
    try:
        from PIL import Image
        file_stream.seek(0)
        img = Image.open(file_stream)
        
        try:
            print("Extracting text from image using Tesseract OCR...")
            ocr_text = pytesseract.image_to_string(img, lang='lao+eng')
            if ocr_text and len(ocr_text.strip()) > 15:
                print("Image text extracted successfully using Tesseract OCR.")
                return ocr_text.strip()
        except Exception as ocr_err:
            print(f"Tesseract OCR failed on image: {ocr_err}")
            
        print("Tesseract OCR yielded insufficient text, falling back to Gemini Vision...")
        import gemini_helper
        file_stream.seek(0)
        img = Image.open(file_stream)
        client = gemini_helper.get_client()
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=["Please extract and transcribe all the text from this image exactly as it appears. If there is no text, describe the image.", img]
        )
        return response.text
    except Exception as e:
        print(f"Error reading Image: {e}")
        return None

def generate_docx_file(test_data, school=None, subject=None, motto=None, grade=None, watermark=None, exam_no=None, time_limit=None):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Noto Sans Lao'
    font.size = Pt(12)
    
    def set_font(run, size=12, bold=False, italic=False):
        run.font.name = 'Noto Sans Lao'
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic

    # 1. Top Meta Row: Grade & Watermark
    meta_table = doc.add_table(rows=1, cols=2)
    meta_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_table.autofit = False
    
    cell_left = meta_table.rows[0].cells[0]
    cell_right = meta_table.rows[0].cells[1]
    cell_left.width = Inches(3.2)
    cell_right.width = Inches(3.2)
    
    p_left = cell_left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_left = p_left.add_run(f"ຊັ້ນ: {grade or 'ມ.7'}")
    set_font(run_left, 11, bold=True)
    

    doc.add_paragraph()
    
    # 2. Lao Motto
    if not motto:
        motto = (
            "ສາທາລະນະລັດ ປະຊາທິປະໄຕ ປະຊາຊົນລາວ\n"
            "ສັນຕິພາບ ເອກະລາດ ປະຊາທິປະໄຕ ເອກະພາບ ວັດທະນາຖາວອນ\n"
            "------000-------"
        )
    for line in motto.split('\n'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line.strip())
        set_font(r, 11, bold=True)
        
    doc.add_paragraph()
    
    # 3. School & Exam No Row
    school_table = doc.add_table(rows=1, cols=2)
    school_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    s_left = school_table.rows[0].cells[0]
    s_right = school_table.rows[0].cells[1]
    s_left.width = Inches(4.5)
    s_right.width = Inches(2.0)
    
    p_school = s_left.paragraphs[0]
    p_school.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_school = p_school.add_run(f"ໂຮງຮຽນ: {school or '........................'}")
    set_font(r_school, 11)
    
    p_exam = s_right.paragraphs[0]
    p_exam.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_exam = p_exam.add_run(f"ເລກທີ: {exam_no or '........'}")
    set_font(r_exam, 11)
    
    # 4. Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(6)
    title_run = title_p.add_run(f"ຫົວບົດສອບເສັງ: {test_data['title']}")
    set_font(title_run, 15, bold=True)
    
    # 5. Subject & Time Row
    sub_table = doc.add_table(rows=1, cols=2)
    sub_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub_left = sub_table.rows[0].cells[0]
    sub_right = sub_table.rows[0].cells[1]
    sub_left.width = Inches(4.5)
    sub_right.width = Inches(2.0)
    
    p_sub = sub_left.paragraphs[0]
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub_str = subject or 'ບົດຮຽນທົ່ວໄປ'
    if not sub_str.startswith('ວິຊາ'):
        sub_str = f"ວິຊາ: {sub_str}"
    r_sub = p_sub.add_run(sub_str)
    set_font(r_sub, 11, bold=True)
    
    p_time = sub_right.paragraphs[0]
    p_time.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_time = p_time.add_run(f"ເວລາ: {time_limit or '90'} ນາທີ")
    set_font(r_time, 11)
    
    doc.add_paragraph()
    
    # 6. Student Grid Table
    grid_table = doc.add_table(rows=1, cols=4)
    grid_table.style = 'Table Grid'
    grid_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    g_cells = grid_table.rows[0].cells
    g_cells[0].width = Inches(3.8)
    g_cells[1].width = Inches(0.9)
    g_cells[2].width = Inches(0.9)
    g_cells[3].width = Inches(0.9)
    
    p_det = g_cells[0].paragraphs[0]
    p_det.paragraph_format.line_spacing = 1.3
    set_font(p_det.add_run("ຊື່ ແລະ ນາມສະກຸນ: ....................................................\nຫ້ອງ: ....................................................\nວັນທີ: ....................................................\nເລກໂຕະ: ...................................................."), 10)
    
    p_sc = g_cells[1].paragraphs[0]
    p_sc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p_sc.add_run("ຄະແນນ"), 10, bold=True)
    
    p_cmt = g_cells[2].paragraphs[0]
    p_cmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p_cmt.add_run("ຄຳເຫັນຂອງຄູ"), 10, bold=True)
    
    p_sgn = g_cells[3].paragraphs[0]
    p_sgn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p_sgn.add_run("ລາຍເຊັນຄູ"), 10, bold=True)
    
    doc.add_paragraph()
    lao_options = {'A': 'ກ', 'B': 'ຂ', 'C': 'ຄ', 'D': 'ງ'}
    
    objective_qs = [q for q in test_data['questions'] if not (q['option_a'] == '' and q['option_b'] == '' and q['option_c'] == '' and q['option_d'] == '')]
    subjective_qs = [q for q in test_data['questions'] if (q['option_a'] == '' and q['option_b'] == '' and q['option_c'] == '' and q['option_d'] == '')]

    if objective_qs:
        sec1_p = doc.add_paragraph()
        sec1_run = sec1_p.add_run("I. ພາກປາລະໄນ (ຄຳຖາມເລືອກຕອບ)")
        sec1_run.bold = True
        sec1_run.font.size = Pt(13)
        doc.add_paragraph()
        
        for idx, q in enumerate(objective_qs, 1):
            qp = doc.add_paragraph()
            q_run = qp.add_run(f"ຂໍ້ {idx}. {q['question_text']}")
            q_run.bold = True
            
            for opt_key, opt_val in [('A', q['option_a']), ('B', q['option_b']), ('C', q['option_c']), ('D', q['option_d'])]:
                if opt_val.strip() != '':
                    op = doc.add_paragraph()
                    op.paragraph_format.left_indent = Inches(0.4)
                    op.paragraph_format.space_after = Pt(2)
                    opt_lao = lao_options[opt_key]
                    op.add_run(f"{opt_lao}. {opt_val}")
            doc.add_paragraph()

    if subjective_qs:
        sec2_p = doc.add_paragraph()
        sec2_run = sec2_p.add_run("II. ພາກອັດຕະໄນ (ຄຳຖາມອະທິບາຍ/ຕອບສັ້ນ)")
        sec2_run.bold = True
        sec2_run.font.size = Pt(13)
        doc.add_paragraph()
        
        for idx, q in enumerate(subjective_qs, 1):
            qp = doc.add_paragraph()
            q_run = qp.add_run(f"ຂໍ້ {idx}. {q['question_text']}")
            q_run.bold = True
            
            for _ in range(3):
                op = doc.add_paragraph()
                op.paragraph_format.left_indent = Inches(0.4)
                op.paragraph_format.space_after = Pt(2)
                op.add_run("....................................................................................................................................................")
            doc.add_paragraph()
        
    doc.add_page_break()
    
    key_title_p = doc.add_paragraph()
    key_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    key_title_run = key_title_p.add_run("ສະເລີຍຄຳຕອບ ແລະ ຄຳອະທິບາຍ (Answer Key)")
    key_title_run.bold = True
    key_title_run.font.size = Pt(14)
    doc.add_paragraph()
    
    for i, q in enumerate(test_data['questions'], 1):
        kp = doc.add_paragraph()
        is_subjective = (q['option_a'] == '' and q['option_b'] == '' and q['option_c'] == '' and q['option_d'] == '')
        
        if is_subjective:
            k_run1 = kp.add_run(f"ຂໍ້ {i}. ແນວທາງຄຳຕອບ: ")
            k_run2 = kp.add_run(q.get('explanation') or 'ບໍ່ມີແນວທາງຄຳຕອບ')
            k_run2.bold = True
        else:
            correct_lao = lao_options.get(q['correct_option'].upper(), q['correct_option'])
            k_run1 = kp.add_run(f"ຂໍ້ {i}. ຕອບ: ")
            k_run2 = kp.add_run(f"{correct_lao}")
            k_run2.bold = True
            
            opt_field = f"option_{q['correct_option'].lower()}"
            correct_text = q.get(opt_field, '')
            kp.add_run(f" ({correct_text})\n")
            
            explanation_text = q.get('explanation') or 'ບໍ່ມີຄຳອະທິບາຍ'
            k_run3 = kp.add_run(f"ອະທິບາຍ: {explanation_text}")
            k_run3.font.size = Pt(10.5)
            k_run3.italic = True
        
        doc.add_paragraph()
        
    docx_stream = io.BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)
    return docx_stream

# Auth Routes
@app.post("/api/auth/register")
async def register(request: Request):
    import re
    try:
        data = await request.json()
    except Exception:
        return JSONResponse(status_code=400, content={"error": "ຮູບແບບຂໍ້ມູນບໍ່ຖືກຕ້ອງ (Invalid JSON)"})
        
    if not data or 'username' not in data or 'password' not in data:
        return JSONResponse(status_code=400, content={"error": "ຂໍ້ມູນບໍ່ຄົບຖ້ວນ"})
    
    # Validate types
    if not isinstance(data['username'], str) or not isinstance(data['password'], str):
        return JSONResponse(status_code=400, content={"error": "ປະເພດຂໍ້ມູນບໍ່ຖືກຕ້ອງ"})
        
    username = data['username'].strip()
    password = data['password']
    
    if not username or not password:
        return JSONResponse(status_code=400, content={"error": "ຊື່ຜູ້ໃຊ້ ຫຼື ລະຫັດຜ່ານບໍ່ສາມາດຫວ່າງເປົ່າໄດ້"})
        
    if len(username) < 3 or len(username) > 30:
        return JSONResponse(status_code=400, content={"error": "ຊື່ຜູ້ໃຊ້ຕ້ອງມີຄວາມຍາວລະຫວ່າງ 3 ຫາ 30 ຕົວອັກສອນ"})
        
    if len(password) < 4 or len(password) > 128:
        return JSONResponse(status_code=400, content={"error": "ລະຫັດຜ່ານຕ້ອງມີຄວາມຍາວລະຫວ່າງ 4 ຫາ 128 ຕົວອັກສອນ"})
        
    # Username character whitelist (alphanumeric, underscore, hyphen)
    if not re.match(r"^[a-zA-Z0-9_-]+$", username):
        return JSONResponse(status_code=400, content={"error": "ຊື່ຜູ້ໃຊ້ຕ້ອງມີສະເພາະຕົວອັກສອນພາສາອັງກິດ, ຕົວເລກ, ຂີດກາງ (-) ແລະ ຂີດກ້ອງ (_) ເທົ່ານັ້ນ"})
        
    try:
        existing_user = database.get_user_by_username(username)
        if existing_user:
            return JSONResponse(status_code=400, content={"error": "ຊື່ຜູ້ໃຊ້ນີ້ຖືກໃຊ້ໄປແລ້ວ"})
            
        password_hash = generate_password_hash(password)
        user_id = database.create_user(username, password_hash)
        if user_id:
            request.session['user_id'] = user_id
            request.session['username'] = username
            request.session['is_guest'] = False
            return {"message": "ລົງທະບຽນສຳເລັດ", "user_id": user_id, "username": username}
        else:
            return JSONResponse(status_code=500, content={"error": "ບໍ່ສາມາດສ້າງບັນຊີໄດ້"})
    except Exception as e:
        print(f"Registration error: {e}")
        return JSONResponse(status_code=500, content={"error": "ເກີດຂໍ້ຜິດພາດໃນການເຊື່ອມຕໍ່ຖານຂໍ້ມູນ"})

@app.post("/api/auth/login")
async def login(request: Request):
    try:
        data = await request.json()
    except Exception:
        return JSONResponse(status_code=400, content={"error": "ຮູບແບບຂໍ້ມູນບໍ່ຖືກຕ້ອງ (Invalid JSON)"})
        
    if not data or 'username' not in data or 'password' not in data:
        return JSONResponse(status_code=400, content={"error": "ຂໍ້ມູນບໍ່ຄົບຖ້ວນ"})
        
    if not isinstance(data['username'], str) or not isinstance(data['password'], str):
        return JSONResponse(status_code=400, content={"error": "ປະເພດຂໍ້ມູນບໍ່ຖືກຕ້ອງ"})
        
    username = data['username'].strip()
    password = data['password']
    
    if not username or not password:
        return JSONResponse(status_code=400, content={"error": "ຊື່ຜູ້ໃຊ້ ຫຼື ລະຫັດຜ່ານບໍ່ສາມາດຫວ່າງເປົ່າໄດ້"})
    
    try:
        user = database.get_user_by_username(username)
        if not user or not check_password_hash(user['password_hash'], password):
            return JSONResponse(status_code=401, content={"error": "ຊື່ຜູ້ໃຊ້ ຫຼື ລະຫັດຜ່ານບໍ່ຖືກຕ້ອງ"})
            
        request.session['user_id'] = user['id']
        request.session['username'] = user['username']
        request.session['is_guest'] = (user['username'] == 'guest')
        profile_pic = user.get('profile_pic') or f"https://ui-avatars.com/api/?name={user['username']}&background=random"
        return {
            "message": "ເຂົ້າລະບົບສຳເລັດ", 
            "user_id": user['id'], 
            "username": user['username'], 
            "is_guest": request.session['is_guest'],
            "profile_pic": profile_pic
        }
    except Exception as e:
        print(f"Login error: {e}")
        return JSONResponse(status_code=500, content={"error": "ເກີດຂໍ້ຜິດພາດໃນລະບົບ"})

@app.post("/api/auth/guest")
async def guest_login(request: Request):
    try:
        username = 'guest'
        user = database.get_user_by_username(username)
        if not user:
            dummy_hash = generate_password_hash('guest_bypass_hash_code_123')
            user_id = database.create_user(username, dummy_hash, token_limit=50000)
            if not user_id:
                user = database.get_user_by_username(username)
                if user:
                    user_id = user['id']
                else:
                    return JSONResponse(status_code=500, content={"error": "ບໍ່ສາມາດສ້າງບັນຊີ Guest ໄດ້"})
        else:
            user_id = user['id']
            
        request.session['user_id'] = user_id
        request.session['username'] = username
        request.session['is_guest'] = True
        return {
            "message": "ເຂົ້າລະບົບໃນຖານະ Guest ສຳເລັດ", 
            "user_id": user_id, 
            "username": username, 
            "is_guest": True
        }
    except Exception as e:
        print(f"Guest login error: {e}")
        return JSONResponse(status_code=500, content={"error": "ເກີດຂໍ້ຜິດພາດໃນລະບົບ"})

@app.post("/api/auth/logout")
async def logout(request: Request):
    try:
        request.session.clear()
        return {"message": "ອອກຈາກລະບົບສຳເລັດ"}
    except Exception as e:
        print(f"Logout error: {e}")
        return JSONResponse(status_code=500, content={"error": "ເກີດຂໍ້ຜິດພາດໃນລະບົບ"})

@app.get("/api/auth/me")
async def get_me(request: Request):
    try:
        user_id = get_current_user(request)
        if user_id:
            user = database.get_user_by_id(user_id)
            if user:
                username = user['username']
                profile_pic = user.get('profile_pic') or f"https://ui-avatars.com/api/?name={username}&background=random"
                return {
                    "logged_in": True,
                    "user_id": user_id,
                    "username": username,
                    "is_guest": request.session.get('is_guest', False),
                    "profile_pic": profile_pic
                }
            else:
                # If session has user_id but user no longer exists in DB, clean the session
                request.session.clear()
        return {"logged_in": False}
    except Exception as e:
        print(f"Get me error: {e}")
        return {"logged_in": False}

# User Profile pic & delete
@app.post("/api/user/profile-pic")
async def upload_profile_pic(request: Request, file: UploadFile = File(...)):
    user_id = get_current_user(request)
    if not user_id:
        return JSONResponse(status_code=401, content={"error": "ກະລຸນາເຂົ້າລະບົບກ່ອນ", "auth_required": True})
    if request.session.get('is_guest'):
        return JSONResponse(status_code=403, content={"error": "ບັນຊີ Guest ບໍ່ສາມາດປ່ຽນຮູບໂປຣໄຟລ໌ໄດ້"})
    
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ['.jpg', '.jpeg', '.png']:
        return JSONResponse(status_code=400, content={"error": "ຮອງຮັບສະເພາະຮູບພາບ (JPG, PNG)"})
    
    import uuid
    filename = f"{user_id}_{uuid.uuid4().hex}{ext}"
    filepath = os.path.join(UPLOAD_AVATAR_FOLDER, filename)
    file_bytes = await file.read()
    with open(filepath, "wb") as f:
        f.write(file_bytes)
        
    pic_url = f"/uploads/avatars/{filename}"
    database.update_user_profile_pic(user_id, pic_url)
    return {"message": "ອັບເດດຮູບໂປຣໄຟລ໌ສຳເລັດ", "profile_pic": pic_url}

@app.delete("/api/user/account")
async def delete_account(request: Request):
    user_id = get_current_user(request)
    if not user_id:
        return JSONResponse(status_code=401, content={"error": "ກະລຸນາເຂົ້າລະບົບກ່ອນ", "auth_required": True})
    if request.session.get('is_guest'):
        return JSONResponse(status_code=403, content={"error": "ບັນຊີ Guest ບໍ່ສາມາດລົບບັນຊີໄດ້"})
    try:
        database.delete_user_account(user_id)
        request.session.clear()
        return {"message": "ລົບບັນຊີສຳເລັດ"}
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})

# Sources Routes
@app.get("/api/sources")
async def get_sources(request: Request):
    user_id = get_current_user(request)
    if not user_id:
        return JSONResponse(status_code=401, content={"error": "ກະລຸນາເຂົ້າລະບົບກ່ອນ", "auth_required": True})
    try:
        sources = database.list_sources(user_id)
        return sources
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})

@app.post("/api/sources/upload/info")
async def upload_source_info(request: Request, file: UploadFile = File(...)):
    user_id = get_current_user(request)
    if not user_id:
        return JSONResponse(status_code=401, content={"error": "ກະລຸນາເຂົ້າລະບົບກ່ອນ", "auth_required": True})
    
    filename_lower = file.filename.lower()
    if not filename_lower.endswith('.pdf'):
        return {"total_pages": 0, "is_pdf": False}
        
    try:
        file_bytes = await file.read()
        file_stream = io.BytesIO(file_bytes)
        reader = pypdf.PdfReader(file_stream)
        total_pages = len(reader.pages)
        return {"total_pages": total_pages, "is_pdf": True}
    except Exception as e:
        print(f"Error reading PDF info: {e}")
        return JSONResponse(status_code=500, content={"error": f"ເກີດຂໍ້ຜິດພາດໃນການອ່ານໄຟລ໌: {str(e)}"})

@app.post("/api/sources/upload")
async def upload_source(
    request: Request,
    file: UploadFile = File(...),
    page_start: str = Form(None),
    page_end: str = Form(None),
    force_ocr: str = Form(None),
    exclude_pages: str = Form(None),
    api_key: str = Form(None)
):
    user_id = get_current_user(request)
    if not user_id:
        return JSONResponse(status_code=401, content={"error": "ກະລຸນາເຂົ້າລະບົບກ່ອນ", "auth_required": True})
        
    filename_lower = file.filename.lower()
    allowed_extensions = ('.pdf', '.docx', '.jpg', '.jpeg', '.png')
    if not filename_lower.endswith(allowed_extensions):
        return JSONResponse(status_code=400, content={"error": "ຮອງຮັບສະເພາະໄຟລ໌ PDF, DOCX ແລະ ຮູບພາບເທົ່ານັ້ນ"})
        
    try:
        file_bytes = await file.read()
        file_size = len(file_bytes)
        file_stream = io.BytesIO(file_bytes)
        
        p_start = int(page_start) if page_start and page_start.strip() != "" else None
        p_end = int(page_end) if page_end and page_end.strip() != "" else None
        force_ocr_bool = force_ocr == 'true' or force_ocr == '1'
        
        p_exclude_set = set()
        if exclude_pages and exclude_pages.strip():
            for part in exclude_pages.split(","):
                try:
                    p_exclude_set.add(int(part.strip()))
                except ValueError:
                    pass
                    
        extracted_text = None
        if filename_lower.endswith('.pdf'):
            extracted_text = extract_text_from_pdf(
                file_stream, 
                page_start=p_start, 
                page_end=p_end, 
                force_ocr=force_ocr_bool, 
                exclude_pages=p_exclude_set
            )
        elif filename_lower.endswith('.docx'):
            extracted_text = extract_text_from_docx(file_stream)
        else:
            extracted_text = extract_text_from_image(file_stream)
            
        if not extracted_text:
            return JSONResponse(status_code=400, content={"error": "ບໍ່ສາມາດອ່ານຂໍ້ຄວາມຈາກໄຟລ໌ໄດ້ ຫຼື ໄຟລ໌ຫວ່າງເປົ່າ"})
            
        metadata = gemini_helper.analyze_document_metadata(extracted_text, api_key=api_key)

        source_id = database.add_source(file.filename, file_size, extracted_text, user_id)
        return {
            "id": source_id,
            "filename": file.filename,
            "file_size": file_size,
            "message": "ອັບໂຫລດ ແລະ ວິເຄາະບົດຮຽນສຳເລັດ",
            "subject": metadata.get("subject", "ບົດຮຽນທົ່ວໄປ"),
            "grade": metadata.get("grade", "ມ.7")
        }
    except Exception as e:
        print(f"Upload error: {e}")
        return JSONResponse(status_code=500, content={"error": f"ເກີດຂໍ້ຜິດພາດໃນການອັບໂຫລດ: {str(e)}"})

@app.delete("/api/sources/{source_id}")
async def delete_source(request: Request, source_id: int):
    user_id = get_current_user(request)
    if not user_id:
        return JSONResponse(status_code=401, content={"error": "ກະລຸນາເຂົ້າລະບົບກ່ອນ", "auth_required": True})
    try:
        source = database.get_source(source_id, user_id)
        if not source:
            return JSONResponse(status_code=404, content={"error": "ບໍ່ພົບບົດຮຽນນີ້"})
        database.delete_source(source_id, user_id)
        return {"message": "ລົບບົດຮຽນສຳເລັດ"}
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})

@app.get("/api/sources/{source_id}/preview")
async def preview_source(request: Request, source_id: int):
    user_id = get_current_user(request)
    if not user_id:
        return JSONResponse(status_code=401, content={"error": "ກະລຸນາເຂົ້າລະບົບກ່ອນ", "auth_required": True})
    try:
        source = database.get_source(source_id, user_id)
        if not source:
            return JSONResponse(status_code=404, content={"error": "ບໍ່ພົບບົດຮຽນນີ້"})
        return {
            "id": source["id"],
            "filename": source["filename"],
            "text_content": source["text_content"]
        }
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})

# Tests Routes
@app.get("/api/tests")
async def get_tests(request: Request):
    user_id = get_current_user(request)
    if not user_id:
        return JSONResponse(status_code=401, content={"error": "ກະລຸນາເຂົ້າລະບົບກ່ອນ", "auth_required": True})
    try:
        tests = database.list_tests(user_id)
        return tests
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})

@app.get("/api/tests/{test_id}")
async def get_test(request: Request, test_id: int):
    user_id = get_current_user(request)
    if not user_id:
        return JSONResponse(status_code=401, content={"error": "ກະລຸນາເຂົ້າລະບົບກ່ອນ", "auth_required": True})
    try:
        test = database.get_test(test_id, user_id)
        if not test:
            return JSONResponse(status_code=404, content={"error": "ບໍ່ພົບການສອບເສັງນີ້"})
        return test
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})

@app.delete("/api/tests/{test_id}")
async def delete_test(request: Request, test_id: int):
    user_id = get_current_user(request)
    if not user_id:
        return JSONResponse(status_code=401, content={"error": "ກະລຸນາເຂົ້າລະບົບກ່ອນ", "auth_required": True})
    try:
        test = database.get_test(test_id, user_id)
        if not test:
            return JSONResponse(status_code=404, content={"error": "ບໍ່ພົບການສອບເສັງນີ້"})
        database.delete_test(test_id, user_id)
        return {"message": "ລົບການສອບເສັງສຳເລັດ"}
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})

@app.put("/api/tests/{test_id}/rich_text")
async def update_test_rich_text(request: Request, test_id: int):
    user_id = get_current_user(request)
    if not user_id:
        return JSONResponse(status_code=401, content={"error": "ກະລຸນາເຂົ້າລະບົບກ່ອນ", "auth_required": True})
    data = await request.json()
    if not data or 'rich_text_content' not in data:
        return JSONResponse(status_code=400, content={"error": "ບໍ່ມີເນື້ອຫາສົ່ງມາ"})
    try:
        success = database.update_test_rich_text(test_id, user_id, data['rich_text_content'])
        if not success:
            return JSONResponse(status_code=404, content={"error": "ບໍ່ພົບການສອບເສັງ ຫຼື ບໍ່ມີສິດ"})
        return {"message": "ບັນທຶກເອກະສານສຳເລັດ"}
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})

@app.get("/api/system-prompts")
async def get_system_prompts(request: Request):
    user_id = get_current_user(request)
    if not user_id:
        return JSONResponse(status_code=401, content={"error": "ກະລຸນາເຂົ້າລະບົບກ່ອນ", "auth_required": True})
    prompt_dir = os.path.join(os.path.dirname(__file__), 'system-prompt-lao')
    prompts = ['default']
    if os.path.exists(prompt_dir):
        for f in os.listdir(prompt_dir):
            if f.endswith('.txt'):
                prompts.append(f[:-4])
    return prompts

@app.post("/api/tests/generate")
async def generate_test(request: Request):
    user_id = get_current_user(request)
    if not user_id:
        return JSONResponse(status_code=401, content={"error": "ກະລຸນາເຂົ້າລະບົບກ່ອນ", "auth_required": True})
    data = await request.json()
    if not data or ('source_id' not in data and 'source_ids' not in data):
        return JSONResponse(status_code=400, content={"error": "ຂໍ້ມູນບໍ່ຄົບຖ້ວນ (ຕ້ອງການ source_id)"})
        
    source_ids = data.get('source_ids', [])
    if not source_ids and 'source_id' in data:
        if isinstance(data['source_id'], list):
            source_ids = data['source_id']
        else:
            source_ids = [data['source_id']]

    question_type = data.get('question_type', 'multiple_choice')
    custom_instructions = data.get('custom_instructions', '')
    num_options = int(data.get('num_options', 4))
    
    num_objective = data.get('num_objective')
    num_subjective = data.get('num_subjective')
    if num_objective is not None:
        num_objective = int(num_objective)
    if num_subjective is not None:
        num_subjective = int(num_subjective)

    if question_type == 'mixed' and num_objective is not None and num_subjective is not None:
        num_questions = num_objective + num_subjective
    else:
        num_questions = int(data.get('num_questions', 10))

    difficulty = data.get('difficulty', 'medium')
    
    system_prompt_name = data.get('system_prompt', 'default')
    if system_prompt_name and system_prompt_name != 'default':
        import re
        if not re.match(r"^[a-zA-Z0-9_-]+$", system_prompt_name):
            return JSONResponse(status_code=400, content={"error": "ຮູບແບບຊື່ System Prompt ບໍ່ຖືກຕ້ອງ"})
        prompt_path = os.path.join(os.path.dirname(__file__), 'system-prompt-lao', f"{system_prompt_name}.txt")
        if os.path.exists(prompt_path):
            with open(prompt_path, 'r', encoding='utf-8') as f:
                sys_content = f.read()
                if custom_instructions:
                    custom_instructions = f"{sys_content}\n\nUser Custom Instructions:\n{custom_instructions}"
                else:
                    custom_instructions = sys_content
    language = data.get('language', 'lao')
    time_limit = int(data.get('time_limit', 0))
    shuffle_options = data.get('shuffle_options', False)
    
    model = data.get('model', 'gemini-2.5-flash')
    api_keys = data.get('api_keys', {})
    if data.get('api_key'):
        api_keys['gemini'] = data.get('api_key')
        
    if not database.is_user_within_token_limit(user_id):
        return JSONResponse(status_code=400, content={"error": "ທ່ານໄດ້ໃຊ້ໂທເຄັນເກີນກຳນົດແລ້ວ (Token limit exceeded)"})

    try:
        combined_text = ""
        valid_source_ids = []
        for sid in source_ids:
            source = database.get_source(sid, user_id)
            if source:
                valid_source_ids.append(sid)
                combined_text += f"--- ຂໍ້ມູນຈາກໄຟລ໌: {source['filename']} ---\n{source['text_content']}\n\n"
                
        if not combined_text:
            return JSONResponse(status_code=404, content={"error": "ບໍ່ພົບບົດຮຽນທີ່ເລືອກ"})
            
        gemini_response, token_count = llm_helper.generate_test_questions(
            model_name=model,
            context_text=combined_text,
            num_questions=num_questions,
            difficulty_lao=difficulty,
            question_type=question_type,
            custom_instructions=custom_instructions,
            num_options=num_options,
            language=language,
            api_keys=api_keys,
            num_objective=num_objective,
            num_subjective=num_subjective
        )
        
        database.increment_token_usage(user_id, token_count)
        
        if shuffle_options and question_type in ['multiple_choice', 'mixed']:
            import random
            for q in gemini_response['questions']:
                options = [
                    ('A', q.get('option_a', '')),
                    ('B', q.get('option_b', '')),
                    ('C', q.get('option_c', '')),
                    ('D', q.get('option_d', '')),
                ]
                correct_key = q['correct_option']
                correct_text = dict(options).get(correct_key, '')
                random.shuffle(options)
                q['option_a'] = options[0][1]
                q['option_b'] = options[1][1]
                q['option_c'] = options[2][1]
                q['option_d'] = options[3][1]
                for new_key_idx, (_, text) in enumerate(options):
                    if text == correct_text:
                        q['correct_option'] = ['A', 'B', 'C', 'D'][new_key_idx]
                        break
                        
        test_id = database.create_test(
            title=gemini_response['title'],
            difficulty=difficulty,
            num_questions=len(gemini_response['questions']),
            source_id=valid_source_ids[0] if valid_source_ids else None,
            questions=gemini_response['questions'],
            user_id=user_id
        )
        
        saved_test = database.get_test(test_id, user_id)
        saved_test['time_limit'] = time_limit
        return saved_test
        
    except ValueError as ve:
        return JSONResponse(status_code=400, content={"error": str(ve)})
    except Exception as e:
        import traceback
        with open('error_log.txt', 'a', encoding='utf-8') as f:
            f.write(f"Generate Test Error: {e}\n{traceback.format_exc()}\n\n")
        print(f"Generate Test Error: {e}")
        return JSONResponse(status_code=500, content={"error": f"ເກີດຂໍ້ຜິດພາດໃນການສ້າງຄຳຖາມ: {str(e)}"})

# Questions Edit/Delete
@app.put("/api/questions/{question_id}")
async def update_question(request: Request, question_id: int):
    user_id = get_current_user(request)
    if not user_id:
        return JSONResponse(status_code=401, content={"error": "ກະລຸນາເຂົ້າລະບົບກ່ອນ", "auth_required": True})
    data = await request.json()
    if not data:
        return JSONResponse(status_code=400, content={"error": "ບໍ່ມີຂໍ້ມູນສົ່ງມາ"})
        
    try:
        result = database.update_question(
            question_id=question_id,
            question_text=data['question_text'],
            option_a=data['option_a'],
            option_b=data['option_b'],
            option_c=data['option_c'],
            option_d=data['option_d'],
            correct_option=data['correct_option'],
            explanation=data.get('explanation', ''),
            user_id=user_id
        )
        if result is False:
            return JSONResponse(status_code=404, content={"error": "ບໍ່ພົບຄຳຖາມ ຫຼື ບໍ່ມີສິດ"})
        return {"message": "ແກ້ໄຂຄຳຖາມສຳເລັດ"}
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})

@app.delete("/api/questions/{question_id}")
async def delete_question(request: Request, question_id: int):
    user_id = get_current_user(request)
    if not user_id:
        return JSONResponse(status_code=401, content={"error": "ກະລຸນາເຂົ້າລະບົບກ່ອນ", "auth_required": True})
    try:
        database.delete_question(question_id, user_id=user_id)
        return {"message": "ລົບຄຳຖາມສຳເລັດ"}
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})

@app.post("/api/questions")
async def add_question(request: Request):
    user_id = get_current_user(request)
    if not user_id:
        return JSONResponse(status_code=401, content={"error": "ກະລຸນາເຂົ້າລະບົບກ່ອນ", "auth_required": True})
    data = await request.json()
    if not data or 'test_id' not in data:
        return JSONResponse(status_code=400, content={"error": "ข้อมูลບໍ່ຄົບຖ້ວນ"})
        
    try:
        question_id = database.add_question(
            test_id=data['test_id'],
            question_text=data['question_text'],
            option_a=data['option_a'],
            option_b=data['option_b'],
            option_c=data['option_c'],
            option_d=data['option_d'],
            correct_option=data['correct_option'],
            explanation=data.get('explanation', ''),
            user_id=user_id
        )
        if question_id is None:
            return JSONResponse(status_code=404, content={"error": "ບໍ່ພົບບົດສອບເສັງ ຫຼື ບໍ່ມີສິດ"})
        return {"id": question_id, "message": "ເພີ່ມຄຳຖາມສຳເລັດ"}
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})

# Stats
@app.get("/api/dashboard/stats")
async def get_stats(request: Request):
    user_id = get_current_user(request)
    if not user_id:
        return JSONResponse(status_code=401, content={"error": "ກະລຸນາເຂົ້າລະບົບກ່ອນ", "auth_required": True})
    try:
        stats = database.get_dashboard_stats(user_id)
        return stats
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})

# Docx Export (FastAPI StreamingResponse)
@app.api_route("/api/tests/{test_id}/export/docx", methods=["GET", "POST"])
async def export_docx(request: Request, test_id: int):
    user_id = get_current_user(request)
    if not user_id:
        return JSONResponse(status_code=401, content={"error": "ກະລຸນາເຂົ້າລະບົບກ່ອນ", "auth_required": True})
    try:
        test = database.get_test(test_id, user_id)
        if not test:
            return JSONResponse(status_code=404, content={"error": "ບໍ່ພົບການສອບເສັງນີ້"})
            
        school = request.query_params.get("school")
        subject = request.query_params.get("subject")
        motto = request.query_params.get("motto")
        grade = request.query_params.get("grade")
        watermark = request.query_params.get("watermark")
        exam_no = request.query_params.get("exam_no")
        time_limit = request.query_params.get("time_limit")

        docx_stream = generate_docx_file(
            test,
            school=school,
            subject=subject,
            motto=motto,
            grade=grade,
            watermark=watermark,
            exam_no=exam_no,
            time_limit=time_limit
        )
        filename = f"test_{test_id}.docx"
        
        headers = {
            'Content-Disposition': f'attachment; filename="{filename}"'
        }
        return StreamingResponse(
            docx_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers
        )
    except Exception as e:
        print(f"Export error: {e}")
        return JSONResponse(status_code=500, content={"error": f"ເກີດຂໍ້ຜິດພາດໃນການດາວໂຫລດ Word: {str(e)}"})

# Chat Route
@app.post("/api/chat")
async def chat_with_source(request: Request):
    user_id = get_current_user(request)
    if not user_id:
        return JSONResponse(status_code=401, content={"error": "ກະລຸນາເຂົ້າລະບົບກ່ອນ", "auth_required": True})
    data = await request.json()
    if not data or 'source_id' not in data or 'message' not in data:
        return JSONResponse(status_code=400, content={"error": "ຂໍ້ມູນບໍ່ຄົບຖ້ວນ"})
        
    source_id = data['source_id']
    new_message = data['message']
    chat_history = data.get('history', [])
    model = data.get('model', 'gemini-2.5-flash')
    api_keys = data.get('api_keys', {})
    if data.get('api_key'):
        api_keys['gemini'] = data.get('api_key')
        
    if not database.is_user_within_token_limit(user_id):
        return JSONResponse(status_code=400, content={"error": "ທ່ານໄດ້ໃຊ້ໂທເຄັນເກີນກຳນົດແລ້ວ (Token limit exceeded)"})

    try:
        source = database.get_source(source_id, user_id)
        if not source:
            return JSONResponse(status_code=404, content={"error": "ບໍ່ພົບບົດຮຽນນີ້"})
            
        response_text, token_count = llm_helper.generate_chat_response(
            model_name=model,
            chat_history=chat_history,
            new_message=new_message,
            context_text=source['text_content'],
            api_keys=api_keys
        )
        database.increment_token_usage(user_id, token_count)
        return {"response": response_text}
    except ValueError as ve:
        return JSONResponse(status_code=400, content={"error": str(ve)})
    except Exception as e:
        print(f"Chat error: {e}")
        return JSONResponse(status_code=500, content={"error": f"ເກີດຂໍ້ຜິດພາດໃນການສົນທະນາ: {str(e)}"})

# Serve Uploads Avatar
@app.get("/uploads/avatars/{filename:path}")
def uploaded_avatar(filename: str):
    file_path = os.path.join(UPLOAD_AVATAR_FOLDER, filename)
    # Prevent directory traversal
    resolved_path = os.path.abspath(file_path)
    if not resolved_path.startswith(os.path.abspath(UPLOAD_AVATAR_FOLDER)):
        return JSONResponse(status_code=403, content={"detail": "Access denied"})
    if os.path.isfile(file_path):
        return FileResponse(file_path)
    return JSONResponse(status_code=404, content={"detail": "Avatar not found"})

# Static Files serving (Vite dist build)
static_dir = os.path.join(os.path.dirname(__file__), 'frontend', 'dist')

@app.get("/")
def serve_index():
    return FileResponse(os.path.join(static_dir, 'index.html'))

@app.get("/{path:path}")
def serve_static_or_index(path: str):
    if path.startswith("api/"):
        return JSONResponse(status_code=404, content={"error": "Not found"})
    file_path = os.path.join(static_dir, path)
    if os.path.isfile(file_path):
        return FileResponse(file_path)
    return FileResponse(os.path.join(static_dir, 'index.html'))

if __name__ == '__main__':
    import uvicorn
    port = int(os.getenv("PORT", 5000))
    uvicorn.run(app, host='0.0.0.0', port=port)
